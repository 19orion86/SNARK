"""FastAPI endpoints для модуля протоколов совещаний.

Предоставляет REST API:
- POST /api/v1/protocols/upload — загрузить аудиофайл
- GET /api/v1/protocols/ — список протоколов
- GET /api/v1/protocols/{id}/celery-status — статус Celery-задачи обработки
- POST /api/v1/protocols/{id}/retry-processing — снова поставить обработку в очередь Celery
- GET /api/v1/protocols/{id} — детали протокола
- PATCH /api/v1/protocols/action-items/{id}/status — обновить статус поручения
"""

from __future__ import annotations

import re
import uuid
from datetime import date
from io import BytesIO
from pathlib import Path
from typing import Annotated
from urllib.parse import quote

import structlog
from celery.result import AsyncResult
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile, status
from fastapi.responses import StreamingResponse
from docx import Document
from sqlalchemy.ext.asyncio import AsyncSession

from src.core.celery_app import celery_app
from src.core.config import settings
from src.core.database import get_async_session
from src.core.security import decrypt_data, encrypt_file
from src.modules.protocols.models import ActionItemStatus, MeetingProtocol, ProtocolStatus
from src.modules.protocols.repository import ProtocolRepository
from src.modules.protocols.schemas import (
    ActionItemStatusEnum,
    ProtocolCeleryRetryResponseSchema,
    ProtocolCeleryStatusResponseSchema,
    ProtocolListResponseSchema,
    ProtocolResponseSchema,
    ProtocolStatusEnum,
    ProtocolUploadResponseSchema,
)
from src.modules.protocols.tasks import process_meeting_audio
from src.modules.protocols.upload_media import (
    ALLOWED_MEETING_UPLOAD_EXTENSIONS,
    prepare_temp_file_for_encryption,
)

logger = structlog.get_logger(__name__)

router = APIRouter(
    prefix="/api/v1/protocols",
    tags=["protocols"],
)

MAX_FILE_SIZE = 2 * 1024 * 1024 * 1024


def _safe_docx_name(title: str, protocol_id: int) -> str:
    sanitized = re.sub(r"[^\w\-. ]+", "_", title, flags=re.UNICODE).strip()
    if not sanitized:
        sanitized = f"protocol_{protocol_id}"
    return f"{sanitized[:80]}.docx"


def _content_disposition_docx(filename: str) -> str:
    ascii_fallback = re.sub(r"[^A-Za-z0-9._-]+", "_", filename).strip("._")
    if not ascii_fallback:
        ascii_fallback = "protocol.docx"
    if not ascii_fallback.lower().endswith(".docx"):
        ascii_fallback += ".docx"
    quoted_utf8_name = quote(filename, safe="")
    return (
        f'attachment; filename="{ascii_fallback}"; '
        f"filename*=UTF-8''{quoted_utf8_name}"
    )


def _build_protocol_docx(
    protocol: MeetingProtocol,
    transcript_text: str | None,
) -> bytes:
    document = Document()
    document.add_heading("Протокол совещания", level=1)
    document.add_paragraph(f"Название: {protocol.title}")
    document.add_paragraph(f"Дата: {protocol.meeting_date}")
    participants = ", ".join(protocol.participants or []) or "Не указаны"
    document.add_paragraph(f"Участники: {participants}")
    document.add_paragraph(f"Статус: {protocol.status.value}")

    if protocol.protocol_text:
        document.add_heading("Готовый протокол", level=2)
        document.add_paragraph(protocol.protocol_text)

    if protocol.agenda:
        document.add_heading("Повестка", level=2)
        for item in protocol.agenda:
            document.add_paragraph(str(item), style="List Bullet")

    if protocol.decisions:
        document.add_heading("Принятые решения", level=2)
        for item in protocol.decisions:
            document.add_paragraph(str(item), style="List Number")

    if protocol.action_items:
        document.add_heading("Поручения", level=2)
        for idx, item in enumerate(protocol.action_items, 1):
            deadline = str(item.deadline) if item.deadline else "не указан"
            document.add_paragraph(
                f"{idx}. {item.text}\n"
                f"Ответственный: {item.assignee}\n"
                f"Срок: {deadline}\n"
                f"Приоритет: {item.priority.value}\n"
                f"Статус: {item.status.value}",
            )

    if protocol.tone_analysis:
        document.add_heading("Анализ тональности", level=2)
        document.add_paragraph(
            f"Оценка: {protocol.tone_analysis.overall_score}/10"
        )
        document.add_paragraph(
            "Соответствие корпоративной культуре: "
            + ("Да" if protocol.tone_analysis.is_compliant else "Нет")
        )
        if protocol.tone_analysis.recommendations:
            document.add_paragraph("Рекомендации:")
            for rec in protocol.tone_analysis.recommendations:
                document.add_paragraph(str(rec), style="List Bullet")

    if transcript_text:
        document.add_heading("Транскрибация", level=2)
        document.add_paragraph(transcript_text)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


@router.post(
    "/upload",
    response_model=ProtocolUploadResponseSchema,
    status_code=status.HTTP_201_CREATED,
    summary="Загрузить аудиофайл совещания",
    description=(
        "Загрузите запись встречи в формате multipart/form-data.\n\n"
        "Как заполнить форму:\n"
        "- `file`: аудио или видео записи (аудио: mp3, wav, ogg, m4a, webm, opus; "
        "видео: mp4, mov, mkv и др. — на сервере извлекается звук в MP3 через ffmpeg)\n"
        "- `title`: название встречи\n"
        "- `meeting_date`: дата встречи в формате `YYYY-MM-DD`\n"
        "- `participants`: список участников через запятую (необязательно)\n"
        "- `source`: источник (`web` или `telegram`)\n\n"
        "После загрузки запускается фоновая обработка (Celery)."
    ),
)
async def upload_meeting_audio(
    file: Annotated[
        UploadFile,
        File(..., description="Аудиофайл совещания"),
    ],
    title: Annotated[
        str,
        Form(..., min_length=3, max_length=500, description="Название встречи"),
    ],
    meeting_date: Annotated[
        date,
        Form(..., description="Дата встречи в формате YYYY-MM-DD"),
    ],
    participants: Annotated[
        str,
        Form(
            description="Участники через запятую",
        ),
    ] = "",
    source: Annotated[
        str,
        Form(description="Источник: web или telegram"),
    ] = "web",
    session: AsyncSession = Depends(get_async_session),
) -> ProtocolUploadResponseSchema:
    """Загрузить аудиофайл совещания для обработки.

    Принимает multipart/form-data с аудиофайлом и метаданными.
    Шифрует файл и запускает Celery-задачу обработки.

    Args:
        file: аудиофайл (MP3, WAV, OGG, M4A, WEBM).
        title: название совещания.
        meeting_date: дата проведения.
        participants: участники через запятую.
        source: источник (web / telegram).
        session: сессия БД.

    Returns:
        ProtocolUploadResponseSchema с ID протокола и задачи Celery.

    Raises:
        HTTPException: 400 при неверном формате файла или размере.
    """
    file_name = file.filename or "audio.wav"
    file_ext = Path(file_name).suffix.lower()

    if file_ext not in ALLOWED_MEETING_UPLOAD_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=(
                f"Неподдерживаемый формат: {file_ext}. "
                f"Допустимые: {', '.join(sorted(ALLOWED_MEETING_UPLOAD_EXTENSIONS))}"
            ),
        )

    meetings_dir = settings.meetings_dir
    meetings_dir.mkdir(parents=True, exist_ok=True)

    file_uid = uuid.uuid4()
    temp_path = meetings_dir / f"temp_{file_uid}{file_ext}"
    to_encrypt: Path | None = None

    try:
        total_size = 0
        chunk_size = 1024 * 1024
        with temp_path.open("wb") as temp_file:
            while True:
                chunk = await file.read(chunk_size)
                if not chunk:
                    break
                total_size += len(chunk)
                if total_size > MAX_FILE_SIZE:
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail=(
                            "Файл слишком большой. "
                            f"Максимум: {MAX_FILE_SIZE // (1024 * 1024)} МБ"
                        ),
                    )
                temp_file.write(chunk)

        if total_size == 0:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Файл пустой",
            )

        try:
            to_encrypt, enc_suffix = prepare_temp_file_for_encryption(
                temp_path,
                file_ext,
            )
            encrypted_path = meetings_dir / f"enc_{file_uid}{enc_suffix}"
            encrypt_file(to_encrypt, encrypted_path)
        except ValueError as exc:
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail=str(exc),
            ) from exc
        except FileNotFoundError as exc:
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail=str(exc),
            ) from exc
        except RuntimeError as exc:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=str(exc),
            ) from exc
    finally:
        if to_encrypt is not None and to_encrypt.exists():
            to_encrypt.unlink(missing_ok=True)
        if temp_path.exists():
            temp_path.unlink(missing_ok=True)

    participants_list = [
        p.strip() for p in participants.split(",") if p.strip()
    ]

    repo = ProtocolRepository(session)
    protocol = await repo.create_protocol(
        title=title,
        meeting_date=meeting_date,
        file_path=str(encrypted_path),
        file_original_name=file_name,
        source=source,
        participants=participants_list,
    )

    task = process_meeting_audio.delay(
        protocol_id=protocol.id,
        encrypted_file_path=str(encrypted_path),
    )
    await repo.set_celery_task_id(protocol.id, task.id)

    logger.info(
        "Аудиофайл загружен через API",
        protocol_id=protocol.id,
        celery_task_id=task.id,
    )

    return ProtocolUploadResponseSchema(
        protocol_id=protocol.id,
        status=ProtocolStatusEnum.UPLOADED,
        celery_task_id=task.id,
        message="Файл загружен. Обработка начата.",
    )


@router.get(
    "/",
    response_model=list[ProtocolListResponseSchema],
    summary="Список протоколов",
)
async def get_protocols_list(
    limit: int = 50,
    offset: int = 0,
    status_filter: ProtocolStatusEnum | None = None,
    session: AsyncSession = Depends(get_async_session),
) -> list[ProtocolListResponseSchema]:
    """Получить список протоколов с пагинацией.

    Args:
        limit: количество записей (макс. 100).
        offset: смещение.
        status_filter: фильтр по статусу.
        session: сессия БД.

    Returns:
        Список ProtocolListResponseSchema.
    """
    limit = min(limit, 100)
    repo = ProtocolRepository(session)

    db_status = None
    if status_filter:
        db_status = ProtocolStatus(status_filter.value)

    protocols = await repo.get_protocols_list(
        limit=limit,
        offset=offset,
        status=db_status,
    )

    return [
        ProtocolListResponseSchema(
            id=p.id,
            title=p.title,
            meeting_date=p.meeting_date,
            status=ProtocolStatusEnum(p.status.value),
            action_items_count=len(p.action_items) if p.action_items else 0,
            overall_tone_score=(
                p.tone_analysis.overall_score if p.tone_analysis else None
            ),
            created_at=p.created_at,
        )
        for p in protocols
    ]


@router.get(
    "/{protocol_id}/celery-status",
    response_model=ProtocolCeleryStatusResponseSchema,
    summary="Статус Celery-задачи обработки",
)
async def get_protocol_celery_status(
    protocol_id: int,
    session: AsyncSession = Depends(get_async_session),
) -> ProtocolCeleryStatusResponseSchema:
    """Показать состояние фоновой задачи Celery для протокола."""
    repo = ProtocolRepository(session)
    protocol = await repo.get_protocol_by_id(protocol_id)

    if not protocol:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Протокол #{protocol_id} не найден",
        )

    if not protocol.celery_task_id:
        return ProtocolCeleryStatusResponseSchema(
            protocol_id=protocol_id,
            protocol_status=protocol.status.value,
            celery_task_id=None,
            celery_state=None,
            detail="Для протокола не задан celery_task_id",
        )

    async_result = AsyncResult(protocol.celery_task_id, app=celery_app)
    info: str | dict | None
    if async_result.failed():
        err = async_result.result
        info = str(err) if err is not None else "failure"
    elif async_result.successful():
        info = async_result.result if isinstance(async_result.result, dict) else None
    else:
        info = async_result.info if isinstance(async_result.info, dict) else None

    return ProtocolCeleryStatusResponseSchema(
        protocol_id=protocol_id,
        protocol_status=protocol.status.value,
        celery_task_id=protocol.celery_task_id,
        celery_state=async_result.state,
        celery_ready=async_result.ready(),
        celery_info=info,
    )


@router.post(
    "/{protocol_id}/retry-processing",
    response_model=ProtocolCeleryRetryResponseSchema,
    status_code=status.HTTP_200_OK,
    summary="Снова поставить обработку в очередь Celery",
    description=(
        "Используйте, если `celery-status` долго в `PENDING` без движения "
        "(например, после падения воркера до исправления пула на Windows). "
        "Доступно только для статусов `uploaded` и `failed`, пока файл есть на диске."
    ),
)
async def retry_protocol_processing(
    protocol_id: int,
    session: AsyncSession = Depends(get_async_session),
) -> ProtocolCeleryRetryResponseSchema:
    repo = ProtocolRepository(session)
    protocol = await repo.get_protocol_by_id(protocol_id)

    if not protocol:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Протокол #{protocol_id} не найден",
        )

    if protocol.status not in (
        ProtocolStatus.UPLOADED,
        ProtocolStatus.FAILED,
    ):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=(
                "Повторная постановка доступна только для статусов "
                "`uploaded` и `failed`."
            ),
        )

    audio_path = Path(protocol.file_path)
    if not audio_path.is_file():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Зашифрованный аудиофайл не найден на диске.",
        )

    previous = protocol.status.value
    task = process_meeting_audio.delay(
        protocol_id=protocol.id,
        encrypted_file_path=str(audio_path),
    )
    await repo.set_celery_task_id(protocol.id, task.id)

    logger.info(
        "Повторная постановка обработки протокола",
        protocol_id=protocol.id,
        celery_task_id=task.id,
    )

    return ProtocolCeleryRetryResponseSchema(
        protocol_id=protocol.id,
        previous_status=previous,
        celery_task_id=task.id,
    )


@router.patch(
    "/action-items/{item_id}/status",
    summary="Обновить статус поручения",
)
async def update_action_item_status(
    item_id: int,
    new_status: ActionItemStatusEnum,
    session: AsyncSession = Depends(get_async_session),
) -> dict:
    """Обновить статус поручения (action item).

    Args:
        item_id: идентификатор поручения.
        new_status: новый статус.
        session: сессия БД.

    Returns:
        Подтверждение обновления.

    Raises:
        HTTPException: 404 если поручение не найдено.
    """
    repo = ProtocolRepository(session)
    item = await repo.update_action_item_status(
        item_id=item_id,
        status=ActionItemStatus(new_status.value),
    )

    if not item:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Поручение #{item_id} не найдено",
        )

    logger.info(
        "Статус поручения обновлён через API",
        item_id=item_id,
        new_status=new_status.value,
    )

    return {
        "id": item.id,
        "status": item.status.value,
        "updated_at": item.updated_at.isoformat(),
    }


@router.get(
    "/{protocol_id}",
    response_model=ProtocolResponseSchema,
    summary="Детали протокола",
)
async def get_protocol_detail(
    protocol_id: int,
    session: AsyncSession = Depends(get_async_session),
) -> ProtocolResponseSchema:
    """Получить полные данные протокола по ID.

    Args:
        protocol_id: идентификатор протокола.
        session: сессия БД.

    Returns:
        ProtocolResponseSchema с полными данными.

    Raises:
        HTTPException: 404 если протокол не найден.
    """
    repo = ProtocolRepository(session)
    protocol = await repo.get_protocol_by_id(protocol_id)

    if not protocol:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Протокол #{protocol_id} не найден",
        )

    response = ProtocolResponseSchema.model_validate(protocol)
    if protocol.transcript_encrypted:
        try:
            response.transcript_text = decrypt_data(protocol.transcript_encrypted)
        except Exception:
            logger.warning(
                "Не удалось расшифровать transcript_encrypted",
                protocol_id=protocol_id,
            )
    return response


@router.get(
    "/{protocol_id}/export-docx",
    summary="Скачать протокол в формате DOCX",
)
async def export_protocol_docx(
    protocol_id: int,
    session: AsyncSession = Depends(get_async_session),
) -> StreamingResponse:
    """Скачать готовый протокол совещания в формате DOCX."""
    repo = ProtocolRepository(session)
    protocol = await repo.get_protocol_by_id(protocol_id)
    if not protocol:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Протокол #{protocol_id} не найден",
        )

    transcript_text: str | None = None
    if protocol.transcript_encrypted:
        try:
            transcript_text = decrypt_data(protocol.transcript_encrypted)
        except Exception:
            logger.warning(
                "Не удалось расшифровать transcript_encrypted для DOCX",
                protocol_id=protocol_id,
            )

    docx_bytes = _build_protocol_docx(protocol, transcript_text)
    filename = _safe_docx_name(protocol.title, protocol_id)
    headers = {
        "Content-Disposition": _content_disposition_docx(filename),
    }
    return StreamingResponse(
        content=BytesIO(docx_bytes),
        media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        headers=headers,
    )
